import os
import uuid
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from langchain_core.tools import tool


# --------------------------------------------------
# Apply Global Style
# --------------------------------------------------

def apply_style(doc):

    section = doc.sections[0]

    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


    normal = doc.styles["Normal"]

    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08



# --------------------------------------------------
# Header
# --------------------------------------------------

def create_header(doc):


    title = doc.add_paragraph()

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)


    run = title.add_run(
        "AUTONOMOUS AI GENERATED DOCUMENT"
    )

    run.bold = True
    run.font.size = Pt(18)



    subtitle = doc.add_paragraph()

    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle.paragraph_format.space_after = Pt(2)


    sub = subtitle.add_run(
        "Created by AI Planning Agent"
    )

    sub.italic = True
    sub.font.size = Pt(11)



    date = doc.add_paragraph()

    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_after = Pt(8)


    date.add_run(
        datetime.now()
        .strftime("%d %B %Y")
    )


    separator = doc.add_paragraph()

    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER

    separator.paragraph_format.space_after = Pt(8)


    separator.add_run(
        "________________________________________"
    )



# --------------------------------------------------
# Heading
# --------------------------------------------------

def add_heading(doc, text):


    paragraph = doc.add_paragraph()


    # important: creates gap between Day1 Day2 etc.

    paragraph.paragraph_format.space_before = Pt(14)

    paragraph.paragraph_format.space_after = Pt(6)



    run = paragraph.add_run(
        text
    )


    run.bold = True

    run.font.size = Pt(14)




# --------------------------------------------------
# Normal Text
# --------------------------------------------------

def add_text(doc,text):


    paragraph = doc.add_paragraph()


    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY
    )


    paragraph.paragraph_format.space_after = Pt(5)



    run = paragraph.add_run(
        text
    )


    run.font.size = Pt(11)




# --------------------------------------------------
# Bullet Points
# --------------------------------------------------

def add_bullet(doc,text):


    bullet = doc.add_paragraph(
        style="List Bullet"
    )


    bullet.paragraph_format.left_indent = Inches(0.25)

    bullet.paragraph_format.space_after = Pt(3)



    run = bullet.add_run(
        text
    )


    run.font.name = "Calibri"

    run.font.size = Pt(11)




# --------------------------------------------------
# Content Formatter
# --------------------------------------------------

def write_content(doc,content):


    heading_words = [

        # General

        "summary",
        "overview",
        "objective",
        "goal",
        "plan",
        "strategy",
        "schedule",
        "timeline",
        "roadmap",
        "steps",
        "analysis",
        "implementation",
        "recommendation",
        "conclusion",
        "checklist",
        "resources",
        "requirements",


        # Universal planning

        "day",
        "week",
        "month",
        "phase",
        "task",
        "section"

    ]



    for line in content.split("\n"):


        line = line.strip()


        if not line:

            continue



        # remove markdown

        line = (
            line
            .replace("###","")
            .replace("##","")
            .replace("**","")
            .strip()
        )



        lower = line.lower()



        # Detect headings

        is_heading = (

            (
                any(

                    word in lower

                    for word in heading_words

                )

                and len(
                    line.split()
                ) <= 10
            )


            or lower.startswith(

                (

                "day ",
                "week ",
                "month ",
                "phase ",
                "task "

                )

            )

        )



        if is_heading:


            add_heading(
                doc,
                line
            )



        elif line.startswith(

            (

            "-",
            "•",
            "*"

            )

        ):


            clean = (

                line
                .replace("-", "",1)
                .replace("•","",1)
                .replace("*","",1)
                .strip()

            )


            add_bullet(
                doc,
                clean
            )



        else:


            add_text(
                doc,
                line
            )



# --------------------------------------------------
# Footer
# --------------------------------------------------

def add_footer(doc):


    footer = (

        doc.sections[0]
        .footer
        .paragraphs[0]

    )


    footer.text = (

        "Generated using Autonomous AI Document Agent"

    )


    footer.alignment = (

        WD_ALIGN_PARAGRAPH.CENTER

    )




# --------------------------------------------------
# LangChain Tool
# --------------------------------------------------

@tool
def create_word_document(
        content:str
)->str:


    """
    Creates a professional formatted DOCX
    document from autonomous agent output.
    """


    os.makedirs(
        "output",
        exist_ok=True
    )


    path = (

        f"output/"
        f"{uuid.uuid4()}.docx"

    )



    doc = Document()


    apply_style(
        doc
    )


    create_header(
        doc
    )


    write_content(
        doc,
        content
    )


    add_footer(
        doc
    )


    doc.save(
        path
    )


    return path